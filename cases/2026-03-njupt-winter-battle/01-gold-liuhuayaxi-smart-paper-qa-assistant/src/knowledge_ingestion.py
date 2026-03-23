"""Document loading, metadata generation, chunking, and indexing services."""

from __future__ import annotations

import asyncio
import re
from pathlib import Path
from typing import TYPE_CHECKING, Awaitable, Callable, Iterable, Literal

from .config import AppConfig
from .models import DocumentRecord, SourceDocument
from .app_utils import build_cache_key, build_file_signature, detect_language, new_doc_id

if TYPE_CHECKING:
    from .retrieval_engine import VectorStoreService


ProgressCallback = Callable[[str], Awaitable[None] | None]


class DocumentIndexer:
    """Turn local files into chunked documents and push them into the vector store."""

    def __init__(self, config: AppConfig, vector_store: VectorStoreService) -> None:
        self.config = config
        self.vector_store = vector_store

    async def ingest_paths(
        self,
        course_id: str,
        file_paths: list[Path],
        source_type: Literal["lecture", "assignment", "paper"],
        rebuild_index: bool = False,
        chunk_size: int | None = None,
        chunk_overlap: int | None = None,
        merge_small_chunks: bool = False,
        min_chunk_size: int | None = None,
        progress_callback: ProgressCallback | None = None,
    ) -> dict[str, int]:
        """Load raw files, split them into chunks concurrently, and persist the chunks.

        Files are processed in a small worker pool so large knowledge bases do
        not spend all their time waiting on one document at a time. Progress
        messages are still emitted per file so the notebook UI can show which
        file is currently being read and split.
        """

        if rebuild_index:
            await _emit_progress(progress_callback, f"正在重建知识库“{course_id}”的索引...")
            await self.vector_store.reset_course(course_id)
        resolved_chunk_size = chunk_size or self.config.chunk_size
        resolved_chunk_overlap = chunk_overlap or self.config.chunk_overlap
        total_files = len(file_paths)
        worker_limit = max(1, min(total_files, max(2, int(self.config.batch_concurrency))))
        semaphore = asyncio.Semaphore(worker_limit)

        async def process_file(index: int, file_path: Path) -> tuple[int, list[SourceDocument], list[SourceDocument]]:
            async with semaphore:
                await _emit_progress(progress_callback, f"正在读取文件 {index}/{total_files}: {file_path.name}")
                loaded_documents = await load_documents(
                    course_id=course_id,
                    file_paths=[file_path],
                    source_type=source_type,
                )
                await _emit_progress(
                    progress_callback,
                    f"正在切片文件 {index}/{total_files}: {file_path.name}（原始片段 {len(loaded_documents)} 个）",
                )
                current_chunks = await asyncio.to_thread(
                    split_documents,
                    loaded_documents,
                    resolved_chunk_size,
                    resolved_chunk_overlap,
                    merge_small_chunks,
                    min_chunk_size,
                )
                await _emit_progress(
                    progress_callback,
                    f"已完成文件 {index}/{total_files}: {file_path.name}，生成 {len(current_chunks)} 个切片",
                )
                return index, loaded_documents, current_chunks

        results = await asyncio.gather(
            *(process_file(index, file_path) for index, file_path in enumerate(file_paths, start=1))
        )
        raw_documents: list[SourceDocument] = []
        chunked_documents: list[SourceDocument] = []
        for _, loaded_documents, current_chunks in sorted(results, key=lambda item: item[0]):
            raw_documents.extend(loaded_documents)
            chunked_documents.extend(current_chunks)
        await _emit_progress(progress_callback, f"正在写入向量库，共 {len(chunked_documents)} 个切片...")
        stored = await self.vector_store.upsert_documents(chunked_documents, progress_callback=progress_callback)
        await _emit_progress(progress_callback, f"向量库写入完成，成功存储 {stored} 个切片。")
        return {
            "raw_documents": len(raw_documents),
            "chunks": len(chunked_documents),
            "stored": stored,
        }

    async def ingest_documents(
        self,
        documents: list[SourceDocument],
        rebuild_index: bool = False,
        chunk_size: int | None = None,
        chunk_overlap: int | None = None,
        merge_small_chunks: bool = False,
        min_chunk_size: int | None = None,
        progress_callback: ProgressCallback | None = None,
    ) -> dict[str, int]:
        """Index already-loaded documents without reading from disk again."""

        if rebuild_index and documents:
            await _emit_progress(
                progress_callback,
                f"正在重建知识库“{documents[0].metadata.get('course_id', '')}”的索引...",
            )
            await self.vector_store.reset_course(str(documents[0].metadata.get("course_id", "")))
        await _emit_progress(progress_callback, "正在切片内存中的文档...")
        chunked_documents = split_documents(
            documents,
            chunk_size=chunk_size or self.config.chunk_size,
            chunk_overlap=chunk_overlap or self.config.chunk_overlap,
            merge_small_chunks=merge_small_chunks,
            min_chunk_size=min_chunk_size,
        )
        await _emit_progress(progress_callback, f"正在写入向量库，共 {len(chunked_documents)} 个切片...")
        stored = await self.vector_store.upsert_documents(chunked_documents, progress_callback=progress_callback)
        await _emit_progress(progress_callback, f"向量库写入完成，成功存储 {stored} 个切片。")
        return {"raw_documents": len(documents), "chunks": len(chunked_documents), "stored": stored}


def build_document_record(
    course_id: str,
    file_path: Path,
    source_type: Literal["lecture", "assignment", "paper"],
) -> DocumentRecord:
    """Create the initial document metadata before the file is parsed."""

    return DocumentRecord(
        doc_id=new_doc_id(),
        course_id=course_id,
        source_type=source_type,
        file_name=file_path.name,
        file_path=str(file_path.resolve()),
        file_ext=file_path.suffix.lstrip(".").lower(),
        language=detect_language(file_path.stem),
    )


def enrich_language(record: DocumentRecord, text: str) -> DocumentRecord:
    """Refresh document language after reading the real text content."""

    language = detect_language(text)
    if hasattr(record, "model_copy"):
        return record.model_copy(update={"language": language})
    payload = record.model_dump()
    payload["language"] = language
    return DocumentRecord(**payload)


def base_metadata(record: DocumentRecord) -> dict[str, object]:
    """Return the shared metadata attached to every chunk from one source file."""

    return {
        "course_id": record.course_id,
        "doc_id": record.doc_id,
        "file_name": record.file_name,
        "file_path": record.file_path,
        "file_ext": record.file_ext,
        "source_type": record.source_type,
        "language": record.language,
        "file_signature": build_file_signature(record.file_path),
    }


async def load_documents(
    course_id: str,
    file_paths: list[Path],
    source_type: Literal["lecture", "assignment", "paper"],
) -> list[SourceDocument]:
    """Load heterogeneous local files into a unified document list."""

    tasks = [
        asyncio.to_thread(_load_single_file, course_id, Path(file_path), source_type)
        for file_path in file_paths
    ]
    loaded_groups = await asyncio.gather(*tasks)
    documents: list[SourceDocument] = []
    for group in loaded_groups:
        documents.extend(group)
    return documents


def split_documents(
    documents: Iterable[SourceDocument],
    chunk_size: int = 800,
    chunk_overlap: int = 120,
    merge_small_chunks: bool = False,
    min_chunk_size: int | None = None,
) -> list[SourceDocument]:
    """Split raw documents into retrieval chunks and optionally merge tiny chunks."""

    chunks: list[SourceDocument] = []
    chunk_config_signature = build_chunk_config_signature(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        merge_small_chunks=merge_small_chunks,
        min_chunk_size=min_chunk_size,
    )
    resolved_min_chunk_size = (
        _resolve_min_chunk_size(chunk_size, min_chunk_size)
        if merge_small_chunks
        else 0
    )
    for document in documents:
        document_chunks: list[SourceDocument] = []
        file_ext = document.metadata.get("file_ext")
        if file_ext == "md":
            sections = _split_markdown_sections(document)
        elif file_ext in {"txt", "docx"}:
            sections = _split_paragraph_sections(document)
        else:
            sections = [document]
        for section in sections:
            document_chunks.extend(_chunk_document(section, chunk_size, chunk_overlap))
        if merge_small_chunks:
            document_chunks = _merge_small_chunks(
                document_chunks,
                chunk_size=chunk_size,
                min_chunk_size=min_chunk_size,
            )
        else:
            document_chunks = _reindex_chunks(document_chunks)
        document_chunks = _annotate_chunk_settings(
            document_chunks,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            merge_small_chunks=merge_small_chunks,
            resolved_min_chunk_size=resolved_min_chunk_size,
            chunk_config_signature=chunk_config_signature,
        )
        chunks.extend(document_chunks)
    return chunks


def build_chunk_config_signature(
    *,
    chunk_size: int,
    chunk_overlap: int,
    merge_small_chunks: bool,
    min_chunk_size: int | None,
) -> str:
    """Build one stable signature for the current chunking parameters."""

    return build_cache_key(
        {
            "chunk_size": int(chunk_size),
            "chunk_overlap": int(chunk_overlap),
            "merge_small_chunks": bool(merge_small_chunks),
            "min_chunk_size": _resolve_min_chunk_size(chunk_size, min_chunk_size) if merge_small_chunks else 0,
        }
    )


def _load_single_file(
    course_id: str,
    file_path: Path,
    source_type: Literal["lecture", "assignment", "paper"],
) -> list[SourceDocument]:
    ext = file_path.suffix.lower().lstrip(".")
    record = build_document_record(course_id=course_id, file_path=file_path, source_type=source_type)
    loader = {
        "pdf": _load_pdf,
        "md": _load_markdown,
        "txt": _load_text,
        "docx": _load_docx,
    }.get(ext)
    if not loader:
        raise ValueError(f"Unsupported file extension: {ext}")
    return loader(file_path, record)


def _load_pdf(file_path: Path, record: DocumentRecord) -> list[SourceDocument]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf is required to load PDF files.") from exc

    reader = PdfReader(str(file_path))
    documents: list[SourceDocument] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        enriched = enrich_language(record, text)
        metadata = base_metadata(enriched)
        metadata.update({"page": page_index, "page_label": f"p.{page_index}"})
        documents.append(SourceDocument(page_content=text, metadata=metadata))
    return documents


def _load_markdown(file_path: Path, record: DocumentRecord) -> list[SourceDocument]:
    text = file_path.read_text(encoding="utf-8")
    enriched = enrich_language(record, text)
    return [SourceDocument(page_content=text, metadata=base_metadata(enriched))]


def _load_text(file_path: Path, record: DocumentRecord) -> list[SourceDocument]:
    text = file_path.read_text(encoding="utf-8")
    enriched = enrich_language(record, text)
    return [SourceDocument(page_content=text, metadata=base_metadata(enriched))]


def _load_docx(file_path: Path, record: DocumentRecord) -> list[SourceDocument]:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("python-docx is required to load DOCX files.") from exc

    doc = DocxDocument(str(file_path))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    enriched = enrich_language(record, text)
    return [SourceDocument(page_content=text, metadata=base_metadata(enriched))]


def _split_markdown_sections(document: SourceDocument) -> list[SourceDocument]:
    lines = document.page_content.splitlines()
    sections: list[SourceDocument] = []
    current_heading = "Introduction"
    buffer: list[str] = []
    for line in lines:
        match = re.match(r"^(#+)\s+(.*)$", line.strip())
        if match:
            if buffer:
                sections.append(
                    SourceDocument(
                        page_content="\n".join(buffer).strip(),
                        metadata={**document.metadata, "section": current_heading, "section_label": current_heading},
                    )
                )
                buffer = []
            current_heading = match.group(2).strip()
        buffer.append(line)
    if buffer:
        sections.append(
            SourceDocument(
                page_content="\n".join(buffer).strip(),
                metadata={**document.metadata, "section": current_heading, "section_label": current_heading},
            )
        )
    return sections or [document]


def _split_paragraph_sections(document: SourceDocument) -> list[SourceDocument]:
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", document.page_content) if part.strip()]
    if not paragraphs:
        return [document]
    return [
        SourceDocument(
            page_content=paragraph,
            metadata={**document.metadata, "section": f"paragraph_{index}", "section_label": f"paragraph_{index}"},
        )
        for index, paragraph in enumerate(paragraphs, start=1)
    ]


def _chunk_document(document: SourceDocument, chunk_size: int, chunk_overlap: int) -> list[SourceDocument]:
    text = document.page_content.strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [_attach_chunk_id(document, index=1)]

    chunks: list[SourceDocument] = []
    start = 0
    index = 1
    while start < len(text):
        end = min(len(text), start + chunk_size)
        candidate = text[start:end]
        if end < len(text):
            boundary = max(
                candidate.rfind("\n\n"),
                candidate.rfind("\n"),
                candidate.rfind("。"),
                candidate.rfind(". "),
                candidate.rfind(" "),
            )
            if boundary > chunk_size // 3:
                end = start + boundary + 1
                candidate = text[start:end]
        chunk_text = candidate.strip()
        if chunk_text:
            chunks.append(
                _attach_chunk_id(
                    SourceDocument(page_content=chunk_text, metadata=document.metadata.copy()),
                    index=index,
                )
            )
        if end >= len(text):
            break
        start = max(0, end - chunk_overlap)
        index += 1
    return chunks


def _attach_chunk_id(document: SourceDocument, index: int) -> SourceDocument:
    metadata = document.metadata.copy()
    metadata.setdefault("chunk_id", new_doc_id(prefix="chunk"))
    metadata["chunk_index"] = index
    metadata.setdefault("merged_from_count", 1)
    if metadata.get("page") and not metadata.get("page_label"):
        metadata["page_label"] = f"p.{metadata['page']}"
    if metadata.get("section") and not metadata.get("section_label"):
        metadata["section_label"] = str(metadata["section"])
    return SourceDocument(page_content=document.page_content, metadata=metadata)


def _annotate_chunk_settings(
    chunks: list[SourceDocument],
    *,
    chunk_size: int,
    chunk_overlap: int,
    merge_small_chunks: bool,
    resolved_min_chunk_size: int,
    chunk_config_signature: str,
) -> list[SourceDocument]:
    annotated: list[SourceDocument] = []
    for chunk in chunks:
        metadata = chunk.metadata.copy()
        metadata["chunk_size"] = int(chunk_size)
        metadata["chunk_overlap"] = int(chunk_overlap)
        metadata["merge_small_chunks"] = bool(merge_small_chunks)
        metadata["min_chunk_size"] = int(resolved_min_chunk_size)
        metadata["chunk_config_signature"] = chunk_config_signature
        annotated.append(SourceDocument(page_content=chunk.page_content, metadata=metadata))
    return annotated


def _merge_small_chunks(
    chunks: list[SourceDocument],
    chunk_size: int,
    min_chunk_size: int | None,
) -> list[SourceDocument]:
    if len(chunks) <= 1:
        return _reindex_chunks(chunks)

    threshold = _resolve_min_chunk_size(chunk_size, min_chunk_size)
    merged: list[SourceDocument] = []
    pending: list[SourceDocument] = []
    pending_length = 0

    def flush_pending() -> None:
        nonlocal pending, pending_length
        if not pending:
            return
        merged.append(_merge_chunk_group(pending) if len(pending) > 1 else pending[0])
        pending = []
        pending_length = 0

    for chunk in chunks:
        chunk_length = len(chunk.page_content)
        if not pending:
            pending = [chunk]
            pending_length = chunk_length
            continue

        separator = 2
        fits = pending_length + separator + chunk_length <= chunk_size
        should_merge = fits and (pending_length < threshold or chunk_length < threshold)
        if should_merge:
            pending.append(chunk)
            pending_length += separator + chunk_length
            continue
        flush_pending()
        pending = [chunk]
        pending_length = chunk_length

    if pending:
        if merged and pending_length < threshold:
            previous = merged[-1]
            if len(previous.page_content) + 2 + pending_length <= chunk_size:
                merged[-1] = _merge_chunk_group([previous] + pending)
            else:
                flush_pending()
        else:
            flush_pending()
    return _reindex_chunks(merged)


def _merge_chunk_group(chunks: list[SourceDocument]) -> SourceDocument:
    first = chunks[0]
    metadata = first.metadata.copy()
    page_labels = [str(item.metadata.get("page_label", "")).strip() for item in chunks if item.metadata.get("page_label")]
    section_labels = [str(item.metadata.get("section_label", "")).strip() for item in chunks if item.metadata.get("section_label")]
    merged_from_count = sum(int(item.metadata.get("merged_from_count", 1)) for item in chunks)
    metadata["chunk_id"] = new_doc_id(prefix="chunk")
    metadata["merged_from_count"] = merged_from_count
    if page_labels:
        metadata["page_label"] = _merge_locator_labels(page_labels)
    if section_labels:
        metadata["section_label"] = _merge_locator_labels(section_labels)
        metadata["section"] = metadata["section_label"]
    merged_text = "\n\n".join(item.page_content.strip() for item in chunks if item.page_content.strip())
    return SourceDocument(page_content=merged_text, metadata=metadata)


def _merge_locator_labels(labels: list[str]) -> str:
    if not labels:
        return ""
    if len(set(labels)) == 1:
        return labels[0]
    return f"{labels[0]} -> {labels[-1]}"


def _resolve_min_chunk_size(chunk_size: int, min_chunk_size: int | None) -> int:
    if min_chunk_size is not None:
        return max(1, min_chunk_size)
    return max(200, chunk_size // 2)


def _reindex_chunks(chunks: list[SourceDocument]) -> list[SourceDocument]:
    reindexed: list[SourceDocument] = []
    for index, chunk in enumerate(chunks, start=1):
        metadata = chunk.metadata.copy()
        metadata["chunk_index"] = index
        metadata.setdefault("merged_from_count", 1)
        if metadata.get("page") and not metadata.get("page_label"):
            metadata["page_label"] = f"p.{metadata['page']}"
        if metadata.get("section") and not metadata.get("section_label"):
            metadata["section_label"] = str(metadata["section"])
        reindexed.append(SourceDocument(page_content=chunk.page_content, metadata=metadata))
    return reindexed


async def _emit_progress(callback: ProgressCallback | None, message: str) -> None:
    """Forward ingestion progress messages to an optional notebook callback."""

    if callback is None:
        return
    result = callback(message)
    if asyncio.iscoroutine(result):
        await result
